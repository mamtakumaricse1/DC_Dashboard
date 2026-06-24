"""
Convert all TPI markdown docs in docs/ to formatted Word (.docx) files.
Usage: python docs/md_to_docx.py
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = Path(__file__).parent
OUT_DIR = DOCS_DIR / "docx"

# Skip index; executive summary has its own generator
SKIP = {"README.md"}

PRIMARY = RGBColor(0x1E, 0x3A, 0x5F)
MUTED = RGBColor(0x5C, 0x6B, 0x7A)


def set_doc_defaults(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(10)


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_inline(text, paragraph):
    """Add runs with bold / code styling."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|[^*`]+)")
    for part in pattern.findall(text):
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif part.startswith("[") and "](" in part:
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                r = paragraph.add_run(m.group(1))
                r.font.color.rgb = PRIMARY
                r.underline = True
        else:
            paragraph.add_run(part)


def is_table_row(line):
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def is_table_sep(line):
    s = line.strip().replace(" ", "")
    return bool(re.match(r"^\|?[-:|]+\|[-:|]+\|?$", s))


def parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            val = row[j] if j < len(row) else ""
            val = re.sub(r"\*\*([^*]+)\*\*", r"\1", val)
            val = val.replace("<br>", "\n")
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(9)
            if i == 0:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
    doc.add_paragraph()


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    text = "\n".join(lines)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)


def convert_md_to_docx(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    set_doc_defaults(doc)

    # Cover line
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("Tirap Performance Index (TPI)")
    tr.bold = True
    tr.font.size = Pt(11)
    tr.font.color.rgb = PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(md_path.stem.replace("_", " "))
    sr.font.size = Pt(9)
    sr.font.color.rgb = MUTED
    add_horizontal_line(doc)

    i = 0
    in_code = False
    code_buf = []
    list_buf = []
    list_style = None

    def flush_list():
        nonlocal list_buf, list_style
        for item in list_buf:
            style = "List Number" if list_style == "num" else "List Bullet"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.space_after = Pt(2)
            parse_inline(item, p)
        list_buf = []
        list_style = None

    while i < len(lines):
        line = lines[i]
        raw = line.rstrip()

        if raw.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                flush_list()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        if not raw.strip():
            flush_list()
            i += 1
            continue

        if raw.strip() == "---":
            flush_list()
            add_horizontal_line(doc)
            i += 1
            continue

        # Markdown table
        if is_table_row(raw):
            flush_list()
            table_rows = []
            while i < len(lines) and is_table_row(lines[i]):
                if not is_table_sep(lines[i]):
                    table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.+)$", raw)
        if m:
            flush_list()
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            sizes = {1: 16, 2: 13, 3: 11, 4: 10}
            p = doc.add_paragraph()
            r = p.add_run(heading_text)
            r.bold = True
            r.font.size = Pt(sizes.get(level, 10))
            r.font.color.rgb = PRIMARY
            p.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Bullet / numbered list
        bm = re.match(r"^(\s*)[-*]\s+(.+)$", raw)
        nm = re.match(r"^(\s*)\d+\.\s+(.+)$", raw)
        cm = re.match(r"^(\s*)- \[ \]\s+(.+)$", raw)
        if cm:
            flush_list()
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(f"☐ {cm.group(2)}", p)
            i += 1
            continue
        if bm:
            if list_style not in (None, "bullet"):
                flush_list()
            list_style = "bullet"
            list_buf.append(bm.group(2))
            i += 1
            continue
        if nm:
            if list_style not in (None, "num"):
                flush_list()
            list_style = "num"
            list_buf.append(nm.group(2))
            i += 1
            continue

        flush_list()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        parse_inline(raw, p)
        i += 1

    flush_list()
    if in_code and code_buf:
        add_code_block(doc, code_buf)

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(f"Source: {md_path.name}  |  Tirap District Administration")
    fr.font.size = Pt(8)
    fr.italic = True
    fr.font.color.rgb = MUTED

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    return docx_path


def main():
    md_files = sorted(DOCS_DIR.glob("*.md"))
    created = []
    for md in md_files:
        if md.name in SKIP:
            continue
        out = OUT_DIR / f"{md.stem}.docx"
        convert_md_to_docx(md, out)
        created.append(out)
        print(f"OK  {out.name}")

    # README index
    readme = DOCS_DIR / "README.md"
    if readme.exists():
        out = OUT_DIR / "00_Documentation_Index.docx"
        convert_md_to_docx(readme, out)
        created.append(out)
        print(f"OK  {out.name}")

    # Executive summary (separate 1-pager; skip if file open in Word)
    exec_script = DOCS_DIR / "generate_srs_executive_summary.py"
    exec_out = DOCS_DIR / "SRS_Executive_Summary.docx"
    if exec_script.exists():
        import subprocess
        try:
            subprocess.run(["python", str(exec_script)], check=True)
            created.append(exec_out)
            print("OK  SRS_Executive_Summary.docx")
        except (subprocess.CalledProcessError, PermissionError):
            print("SKIP SRS_Executive_Summary.docx (close file in Word and re-run)")

    print(f"\nCreated {len(created)} DOCX file(s) in:")
    print(f"  {OUT_DIR}")
    print(f"  {DOCS_DIR} (executive summary)")


if __name__ == "__main__":
    main()
