"""Generate 1-page SRS Executive Summary as DOCX."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"D:\DistrictD\docs\SRS_Executive_Summary.docx"


def set_narrow_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_meta_table(doc):
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Document ID", "SRS-TPI-UD-001-EXEC"),
        ("Version / Date", "1.0  |  8 June 2026"),
        ("Prepared by", "User Department, District Administration (Tirap)"),
        ("Full SRS reference", "docs/05_SRS_User_Department.md (SRS-TPI-UD-001)"),
    ]
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(9)
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()


def add_bullets(doc, items, size=9, spacing=1):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(spacing)
        p.paragraph_format.space_before = Pt(0)
        for r in p.runs:
            r.font.size = Pt(size)


def add_heading(doc, text, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def build():
    doc = Document()
    set_narrow_margins(doc)

    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("EXECUTIVE SUMMARY")
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Tirap Performance Index (TPI)\nDistrict Performance Monitoring System")
    sr.font.size = Pt(10)
    sr.bold = True
    sub.paragraph_format.space_after = Pt(4)

    add_horizontal_line(doc)
    add_meta_table(doc)

    add_heading(doc, "1. Purpose")
    doc.add_paragraph(
        "This executive summary supports senior approval of the full Software Requirements "
        "Specification (SRS) prepared from the User Department perspective. TPI is a web-based "
        "system for monthly monitoring of 14 Key Result Areas (KRAs), ~100 KPIs, and District "
        "Commissioner (DC) follow-up on underperformance."
    ).paragraph_format.space_after = Pt(4)
    for r in doc.paragraphs[-1].runs:
        r.font.size = Pt(9)

    add_heading(doc, "2. Problem & Solution")
    add_bullets(doc, [
        "Problem: DC review data is scattered — who is RED, who has not submitted, and which targets are overdue is unclear.",
        "Solution: One dashboard (Command Center) plus submission tracker, target follow-up, action tracker, history, and meeting exports.",
    ])

    add_heading(doc, "3. Users")
    t = doc.add_table(rows=4, cols=3)
    t.style = "Table Grid"
    hdr = ["User", "Role", "Primary use"]
    data = [
        ("District Commissioner", "ADMIN", "Daily alerts, monthly review, set targets, export"),
        ("14 KRA Nodal Officers (HoD)", "DEPT", "Submit previous month KPI actuals once per cycle"),
        ("District IT / DIO", "Support", "Deploy, backup SQL Server, manage users"),
    ]
    for j, h in enumerate(hdr):
        t.rows[0].cells[j].text = h
        t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
            for p in t.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    doc.add_paragraph()

    add_heading(doc, "4. Core Capabilities (Must-Have)")
    add_bullets(doc, [
        "Command Center — TPI gauge, 14-KRA heatmap, pending submissions, RED departments, overdue targets",
        "Submission tracker — 14 departments × submitted / pending / late",
        "Target follow-up — DC commitments checked against current month (no manual month switch)",
        "Action tracker — set targets for RED KPIs; review prior commitments",
        "KRA drill-down — all KPIs with actual vs target",
        "History — 3/6 month trends, all departments, distinct colours per KRA",
        "Export — Excel (CSV) and PDF (print-ready HTML) for review meetings",
    ], spacing=0)

    add_heading(doc, "5. Reporting Cycle & RAG Rules")
    p = doc.add_paragraph()
    p.add_run("Cycle: ").bold = True
    p.add_run(
        "Departments submit previous calendar month only (e.g. in June → May data). "
        "DC reviews previous month by default. Target follow-up uses current month scores."
    )
    for r in p.runs:
        r.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(3)

    rag = doc.add_table(rows=4, cols=3)
    rag.style = "Table Grid"
    for j, h in enumerate(["RAG Status", "Score (0–100)", "Action"]):
        rag.rows[0].cells[j].text = h
        rag.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, row in enumerate([
        ("GREEN", "≥ 90", "On track"),
        ("YELLOW", "70 – 89", "Watch list"),
        ("RED", "< 70", "Mandatory DC action plan"),
    ], 1):
        for j, val in enumerate(row):
            rag.rows[i].cells[j].text = val
            for par in rag.rows[i].cells[j].paragraphs:
                for r in par.runs:
                    r.font.size = Pt(8)
    doc.add_paragraph()

    add_heading(doc, "6. Deployment & Scope")
    add_bullets(doc, [
        "Runs on district laptop or office LAN — SQL Server + Node API + browser; cloud hosting not required.",
        "Multi-district ready — separate database and DistrictConfig per district.",
        "Out of scope (v1.0): mobile apps, SMS alerts, automatic pull from external MIS, state-level dashboard.",
    ], spacing=0)

    add_heading(doc, "7. Recommendation")
    p = doc.add_paragraph(
        "User Department recommends approval of SRS-TPI-UD-001 for User Acceptance Testing (UAT) "
        "and operational use by DC office and 14 KRAs. Full SRS contains 60+ functional requirements, "
        "6 use cases, business rules, and sign-off criteria."
    )
    for r in p.runs:
        r.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(6)

    add_heading(doc, "8. Approval (Sign-Off)")
    sig = doc.add_table(rows=5, cols=4)
    sig.style = "Table Grid"
    for j, h in enumerate(["Role", "Name", "Signature", "Date"]):
        sig.rows[0].cells[j].text = h
        sig.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, role in enumerate([
        "District Commissioner",
        "User Department Nodal Officer",
        "District Informatics Officer (IT)",
        "Development / Implementation Lead",
    ], 1):
        sig.rows[i].cells[0].text = role
        for j in range(1, 4):
            sig.rows[i].cells[j].text = ""
        for cell in sig.rows[i].cells:
            for par in cell.paragraphs:
                for r in par.runs:
                    r.font.size = Pt(8)

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        "Tirap Performance Index  |  Executive Summary of SRS-TPI-UD-001  |  User Department, Tirap District"
    )
    fr.font.size = Pt(7)
    fr.italic = True
    fr.font.color.rgb = RGBColor(0x5C, 0x6B, 0x7A)

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()
